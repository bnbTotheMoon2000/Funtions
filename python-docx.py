from docx import Document

def replace_placeholder(doc, placeholders):
    """
    读取Word文档并替换占位符
    :param doc: Document 对象
    :param placeholders: 占位符字典 {占位符: 替换内容}
    """
    for paragraph in doc.paragraphs:
        for placeholder, replacement in placeholders.items():
            if placeholder in paragraph.text:
                # 替换段落中的文本
                paragraph.text = paragraph.text.replace(placeholder, replacement)

    # 处理表格中的占位符
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, replacement in placeholders.items():
                        if placeholder in paragraph.text:
                            # 替换表格中的文本
                            paragraph.text = paragraph.text.replace(placeholder, replacement)

def main():
    # 读取已有的Word文档
    template_path = 'template.docx'  # 样式文档路径
    doc = Document(template_path)

    # 定义占位符和它们的替换内容
    placeholders = {
        '[name]': 'John Doe',
        '[date]': '2024-09-25',
        '[location]': 'New York'
    }

    # 替换占位符
    replace_placeholder(doc, placeholders)
